from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse
from django.contrib import messages
from django.db.models import Q
from io import BytesIO
import xml.etree.ElementTree as ET
import os
from django.urls import reverse
import pandas as pd
import mammoth
from io import BytesIO
import os
from django.http import FileResponse, Http404
from django.shortcuts import get_object_or_404
from .models import Report
import shutil


from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.units import cm

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

from .models import Report, Manager, Category, Product, ReportItem
from .forms import (
    ManagerForm, CategoryForm, ProductForm,
    ReportForm, ReportItemFormSet, SearchForm
)


def admin_panel(request):
    if request.method == 'POST':
        if 'manager_submit' in request.POST:
            form = ManagerForm(request.POST)
            if form.is_valid(): form.save(); messages.success(request, "Менеджер добавлен")
        elif 'category_submit' in request.POST:
            form = CategoryForm(request.POST)
            if form.is_valid(): form.save(); messages.success(request, "Категория добавлена")
        elif 'product_submit' in request.POST:
            form = ProductForm(request.POST)
            if form.is_valid(): form.save(); messages.success(request, "Товар добавлен")
        return redirect('admin_panel')

    return render(request, 'reports/admin_panel.html', {
        'managers': Manager.objects.all(),
        'categories': Category.objects.all(),
        'products': Product.objects.all(),
        'manager_form': ManagerForm(),
        'category_form': CategoryForm(),
        'product_form': ProductForm()
    })


def report_list(request):
    reports = Report.objects.select_related('manager').all()
    return render(request, 'reports/report_list.html', {'reports': reports})


def manager_detail(request, manager_id):
    manager = get_object_or_404(Manager, pk=manager_id)
    reports = Report.objects.filter(manager=manager).order_by('-report_date')
    return render(request, 'reports/manager_detail.html', {'manager': manager, 'reports': reports})


def report_create(request):
    if request.method == 'POST':
        report_form = ReportForm(request.POST)
        item_formset = ReportItemFormSet(request.POST)

        if report_form.is_valid() and item_formset.is_valid():
            report = report_form.save()
            items = item_formset.save(commit=False)
            for item in items:
                item.report = report
                if not item.price_used and item.product:
                    item.price_used = item.product.base_price
                if item.custom_product_name and item.product:
                    item.product = None
                item.save()
            messages.success(request, "Отчёт успешно создан!")
            return redirect('report_list')
    else:
        report_form = ReportForm()
        item_formset = ReportItemFormSet()

    return render(request, 'reports/report_create.html', {
        'report_form': report_form,
        'item_formset': item_formset
    })


def search_reports(request):
    form = SearchForm(request.GET or None)
    reports = Report.objects.select_related('manager').prefetch_related('items').all()

    if form.is_valid():
        cd = form.cleaned_data
        if cd['date_from']: reports = reports.filter(report_date__gte=cd['date_from'])
        if cd['date_to']: reports = reports.filter(report_date__lte=cd['date_to'])
        if cd['manager']: reports = reports.filter(manager=cd['manager'])

        if cd['keyword']:
            kw = cd['keyword']
            q_objects = Q(comments__icontains=kw) | Q(manager__full_name__icontains=kw)
            q_objects |= Q(items__product__name__icontains=kw) | Q(items__custom_product_name__icontains=kw)
            q_objects |= Q(items__category__name__icontains=kw)
            q_objects |= Q(report_date__icontains=kw)  # Поиск по части даты (например "2026" или "-04-")
            q_objects |= Q(items__price_used__icontains=kw)  # Поиск по цене

            reports = reports.filter(q_objects).distinct()

    return render(request, 'reports/search.html', {'form': form, 'reports': reports})


def export_pdf(request, report_id):
    report = get_object_or_404(Report, pk=report_id)

    arial_path = "C:/Windows/Fonts/arial.ttf"
    if os.path.exists(arial_path):
        try:
            pdfmetrics.registerFont(TTFont('Arial', arial_path))
        except:
            pass

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=2 * cm, leftMargin=2 * cm, topMargin=2 * cm,
                            bottomMargin=2 * cm)
    elements = []

    styles = getSampleStyleSheet()
    style_title = ParagraphStyle(name='Title', fontName='Arial', fontSize=16, alignment=1, spaceAfter=30)
    style_normal = ParagraphStyle(name='Normal', fontName='Arial', fontSize=12, leading=14)

    elements.append(Paragraph(f"Отчёт о продажах №{report.id}", style_title))
    elements.append(Paragraph(f"<b>Менеджер:</b> {report.manager.full_name}", style_normal))
    elements.append(Paragraph(f"<b>Дата:</b> {report.report_date}", style_normal))
    if report.comments:
        elements.append(Paragraph(f"<b>Комментарий:</b> {report.comments}", style_normal))

    elements.append(Spacer(1, 0.5 * cm))

    data = [['Товар', 'Категория', 'Кол-во', 'Цена', 'Сумма']]
    for item in report.items.all():
        name = item.custom_product_name or (item.product.name if item.product else "-")
        cat = item.category.name if item.category else "-"
        qty = str(item.quantity)
        price = f"{float(item.price_used or 0):.2f}"
        total = f"{(float(item.price_used or 0) * item.quantity):.2f}"
        data.append([name, cat, qty, price, total])

    table = Table(data, colWidths=[8 * cm, 3 * cm, 2 * cm, 2 * cm, 2 * cm])
    table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Arial'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BACKGROUND', (0, 0), (-1, 0), '#eeeeee'),
        ('GRID', (0, 0), (-1, -1), 1, '#000000'),
        ('ALIGN', (2, 1), (-1, -1), 'RIGHT'),
    ]))

    elements.append(table)
    doc.build(elements)

    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'inline; filename="report_{report.id}.pdf"'
    return response


def export_docx(request, report_id):
    report = get_object_or_404(Report, pk=report_id)
    document = Document()

    heading = document.add_heading(f'Отчёт о продажах №{report.id}', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph()
    p.add_run(f'Менеджер: ').bold = True
    p.add_run(report.manager.full_name)
    p.add_run(f'\nДата: ').bold = True
    p.add_run(str(report.report_date))
    if report.comments:
        p.add_run(f'\nКомментарий: ').bold = True
        p.add_run(report.comments)

    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Товар'
    hdr_cells[1].text = 'Категория'
    hdr_cells[2].text = 'Кол-во'
    hdr_cells[3].text = 'Цена (BYN)'
    hdr_cells[4].text = 'Сумма (BYN)'

    for item in report.items.all():
        row_cells = table.add_row().cells
        row_cells[0].text = item.custom_product_name or (item.product.name if item.product else "-")
        row_cells[1].text = item.category.name if item.category else "-"
        row_cells[2].text = str(item.quantity)
        row_cells[3].text = f"{float(item.price_used or 0):.2f}"
        row_cells[4].text = f"{(float(item.price_used or 0) * item.quantity):.2f}"

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = HttpResponse(buffer.read(),
                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'inline; filename="report_{report.id}.docx"'
    return response


def export_xlsx(request, report_id):
    report = get_object_or_404(Report, pk=report_id)
    wb = Workbook()
    ws = wb.active
    ws.title = f"Отчет {report.id}"

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="FF071ff5", end_color="FF071ff5", fill_type="solid")
    center_align = Alignment(horizontal="center", vertical="center")
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'),
                         bottom=Side(style='thin'))

    ws.merge_cells('A1:E1')
    cell_title = ws['A1']
    cell_title.value = f"ОТЧЁТ О ПРОДАЖАХ №{report.id}"
    cell_title.font = Font(bold=True, size=14)
    cell_title.alignment = Alignment(horizontal="center")
    ws.append([])

    ws.append(["Менеджер:", report.manager.full_name])
    ws.append(["Дата отчёта:", report.report_date])
    if report.comments: ws.append(["Комментарий:", report.comments])
    ws.append([])

    headers = ["№", "Наименование товара", "Категория", "Количество", "Цена (BYN)", "Сумма (BYN)"]
    ws.append(headers)

    for cell in ws[ws.max_row]:
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align

    total_sum = 0
    for idx, item in enumerate(report.items.all(), 1):
        name = item.custom_product_name or (item.product.name if item.product else "Не указано")
        category = item.category.name if item.category else "-"
        qty = item.quantity
        price = float(item.price_used or 0)
        item_sum = qty * price
        total_sum += item_sum

        ws.append([idx, name, category, qty, price, item_sum])
        for cell in ws[ws.max_row]:
            cell.border = thin_border
            if isinstance(cell.value, (int, float)):
                cell.number_format = '#,##0.00'
                cell.alignment = Alignment(horizontal="right")

    ws.append(["", "", "", "ИТОГО:", "", total_sum])
    total_row = ws[ws.max_row]
    total_row[3].font = Font(bold=True)
    total_row[5].font = Font(bold=True)
    total_row[5].number_format = '#,##0.00'

    column_widths = [5, 40, 20, 10, 15, 15]
    for i, width in enumerate(column_widths, 1):
        ws.column_dimensions[chr(64 + i)].width = width

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'inline; filename=report_{report.id}.xlsx'
    wb.save(response)
    return response


def export_xml(request, report_id):
    report = get_object_or_404(Report, pk=report_id)
    root = ET.Element("Report", id=str(report.id), manager=report.manager.full_name, date=str(report.report_date))

    for item in report.items.all():
        item_el = ET.SubElement(root, "Item")
        name = item.custom_product_name or (item.product.name if item.product else "Неизвестно")
        ET.SubElement(item_el, "Name").text = name
        ET.SubElement(item_el, "Category").text = item.category.name if item.category else "Не указана"
        ET.SubElement(item_el, "Quantity").text = str(item.quantity)
        ET.SubElement(item_el, "Price").text = str(item.price_used or 0)
        total = (item.price_used or 0) * item.quantity
        ET.SubElement(item_el, "Total").text = str(total)

    xml_str = ET.tostring(root, encoding='unicode', method='xml')
    xml_response = '<?xml version="1.0" encoding="UTF-8"?>\n' + xml_str

    response = HttpResponse(xml_response, content_type='text/xml; charset=utf-8')
    response['Content-Disposition'] = f'inline; filename="report_{report.id}.xml"'
    return response


def preview_xlsx_js(request, report_id):
    report = get_object_or_404(Report, pk=report_id)
    file_path = os.path.join('media', 'exports', f'report_{report_id}.xlsx')

    if not os.path.exists(file_path):
        from .views import export_xlsx
        from django.http import HttpResponse
        import tempfile

        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
            response = export_xlsx(request, report_id)
            tmp.write(response.content)
            tmp_path = tmp.name
        shutil.move(tmp_path, file_path)

    if not os.path.exists(file_path):
        raise Http404("Файл не найден")

    return render(request, 'reports/preview_xlsx_js.html',
                  {'report': report, 'file_url': f'/media/exports/report_{report_id}.xlsx'})


def preview_docx_js(request, report_id):
    report = get_object_or_404(Report, pk=report_id)
    file_path = os.path.join('media', 'exports', f'report_{report_id}.docx')

    if not os.path.exists(file_path):
        from .views import export_docx
        import tempfile
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            response = export_docx(request, report_id)
            tmp.write(response.content)
            tmp_path = tmp.name
        shutil.move(tmp_path, file_path)

    if not os.path.exists(file_path):
        raise Http404("Файл не найден")

    return render(request, 'reports/preview_docx_js.html',
                  {'report': report, 'file_url': f'/media/exports/report_{report_id}.docx'})


def get_xlsx_file(request, report_id):
    report = get_object_or_404(Report, pk=report_id)

    from openpyxl import Workbook
    from io import BytesIO

    wb = Workbook()
    ws = wb.active
    ws.title = f"Отчёт_{report.id}"

    headers = ['Товар', 'Категория', 'Кол-во', 'Цена (BYN)', 'Сумма (BYN)']
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num, value=header)
        cell.font = cell.font.copy(bold=True)

    for row_num, item in enumerate(report.items.all(), 2):
        name = item.custom_product_name or (item.product.name if item.product else '—')
        category = item.category.name if item.category else '—'
        qty = item.quantity
        price = float(item.price_used or 0)
        total = qty * price

        ws.cell(row=row_num, column=1, value=name)
        ws.cell(row=row_num, column=2, value=category)
        ws.cell(row=row_num, column=3, value=qty)
        ws.cell(row=row_num, column=4, value=price)
        ws.cell(row=row_num, column=5, value=total)

    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    print(f"XLSX file size: {len(buffer.getvalue())} bytes")

    response = HttpResponse(buffer.read(),
                            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="report_{report_id}.xlsx"'
    return response

def get_docx_file(request, report_id):
    report = get_object_or_404(Report, pk=report_id)

    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    document = Document()
    document.add_heading(f'Отчёт о продажах №{report.id}', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph()
    p.add_run('Менеджер: ').bold = True
    p.add_run(report.manager.full_name)
    p.add_run(f'\nДата: ').bold = True
    p.add_run(str(report.report_date))
    if report.comments:
        p.add_run(f'\nКомментарий: ').bold = True
        p.add_run(report.comments)

    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, title in enumerate(['Товар', 'Категория', 'Кол-во', 'Цена (BYN)', 'Сумма (BYN)']):
        hdr[i].text = title
        hdr[i].paragraphs[0].runs[0].bold = True

    for item in report.items.all():
        row = table.add_row().cells
        row[0].text = item.custom_product_name or (item.product.name if item.product else '—')
        row[1].text = item.category.name if item.category else '—'
        row[2].text = str(item.quantity)
        row[3].text = f"{float(item.price_used or 0):.2f}"
        row[4].text = f"{item.quantity * float(item.price_used or 0):.2f}"

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = HttpResponse(buffer.read(),
                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="report_{report_id}.docx"'
    return response